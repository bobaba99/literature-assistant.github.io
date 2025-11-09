from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
import os
from dotenv import load_dotenv
import openai
from werkzeug.utils import secure_filename
import json
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt
import PyPDF2

# Load environment variables
load_dotenv()

# Initialize Flask app with frontend template and static folders
# Get the frontend path - works both in development and Docker
frontend_path = os.path.join(os.path.dirname(__file__), '../../frontend')
if not os.path.exists(frontend_path):
    # In Docker, frontend is at /app/frontend
    frontend_path = '/app/frontend'

app = Flask(__name__,
            template_folder=frontend_path,
            static_folder=frontend_path,
            static_url_path='')

# CORS configuration for production and development
# Get allowed origins from environment variable, defaulting to localhost for development
allowed_origins = os.getenv('ALLOWED_ORIGINS', 'http://localhost:5001,http://127.0.0.1:5001').split(',')
CORS(app, origins=allowed_origins, supports_credentials=True)

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {'pdf'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

# Ensure directories exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# API Key
openai_api_key = os.getenv('OPENAI_API_KEY')
PORT = int(os.getenv('PORT', 5001))

if not openai_api_key:
    print("Error: OPENAI_API_KEY not found. Please add it to your .env file.")

# Initialize OpenAI client
client_openai = openai.OpenAI(api_key=openai_api_key) if openai_api_key else None


def is_allowed_file(filename):
    """Check if file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(pdf_path):
    """Extract text content from PDF file."""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
    except Exception as error:
        raise Exception(f"Failed to extract text from PDF: {str(error)}")


def load_prompt_template():
    """Load the prompt template from prompt.md file."""
    try:
        prompt_path = os.path.join(os.path.dirname(__file__), '..', '..', 'prompt.md')
        with open(prompt_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as error:
        raise Exception(f"Failed to load prompt template: {str(error)}")


def analyze_with_openai(text_content, prompt_template):
    """Analyze document using OpenAI API."""
    if not client_openai:
        raise Exception("OpenAI API key not configured")

    try:
        # Combine prompt template with document content
        full_prompt = f"{prompt_template}\n\n**INPUT:**\n{text_content}"

        response = client_openai.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are a research methodologist and domain expert specializing in analyzing academic papers."
                },
                {
                    "role": "user",
                    "content": full_prompt
                }
            ]
        )

        return response.choices[0].message.content

    except Exception as error:
        raise Exception(f"OpenAI API error: {str(error)}")




def parse_json_from_response(response_text):
    """Extract and parse JSON from AI response."""
    import re

    # Try to find JSON block in markdown code fence (use greedy matching to get full JSON)
    json_match = re.search(r'```json\s*(\{.*\})\s*```', response_text, re.DOTALL)
    if json_match:
        json_str = json_match.group(1)
    else:
        # Try to find raw JSON by finding balanced braces
        # Find the first { and then count braces to find matching }
        start = response_text.find('{')
        if start == -1:
            return None

        brace_count = 0
        in_string = False
        escape_next = False

        for i in range(start, len(response_text)):
            char = response_text[i]

            # Handle string escapes
            if escape_next:
                escape_next = False
                continue

            if char == '\\':
                escape_next = True
                continue

            # Track if we're inside a string
            if char == '"':
                in_string = not in_string
                continue

            # Only count braces outside of strings
            if not in_string:
                if char == '{':
                    brace_count += 1
                elif char == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        json_str = response_text[start:i+1]
                        break
        else:
            return None

    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        print(f"JSON parsing error: {e}")
        print(f"Attempted to parse: {json_str[:200]}...")
        return None


def safe_str(value):
    """Safely convert any value to string, handling None and special cases."""
    if value is None:
        return "N/A"
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, str):
        return value.strip() if value.strip() else "N/A"
    return str(value)


def format_table(data, headers=None):
    """Format a list of dictionaries as a markdown table."""
    if not data or not isinstance(data, list):
        return ""

    # Extract headers from first item if not provided
    if not headers:
        if isinstance(data[0], dict):
            # Collect all unique keys from all items
            all_keys = []
            for item in data:
                if isinstance(item, dict):
                    for key in item.keys():
                        if key not in all_keys:
                            all_keys.append(key)
            headers = all_keys
        else:
            return ""

    if not headers:
        return ""

    # Create table header - escape pipes in header names
    safe_headers = [str(h).replace("|", "\\|") for h in headers]
    table = "| " + " | ".join(safe_headers) + " |\n"
    table += "| " + " | ".join(["---"] * len(headers)) + " |\n"

    # Add rows
    for item in data:
        if isinstance(item, dict):
            row_values = [safe_str(item.get(h, "")).replace("|", "\\|") for h in headers]
            table += "| " + " | ".join(row_values) + " |\n"

    return table + "\n"


def format_nested_dict(data, indent_level=0):
    """Recursively format nested dictionaries with proper indentation."""
    if data is None:
        return "N/A\n"

    if isinstance(data, str):
        return f"{data}\n\n"

    if isinstance(data, (int, float, bool)):
        return f"{safe_str(data)}\n\n"

    if isinstance(data, list):
        # Check if it's a list of dictionaries (potential table)
        if data and isinstance(data[0], dict):
            # Try table format if we have multiple items with similar structure
            if len(data) > 1:
                first_keys = set(data[0].keys())
                # Check if at least 50% of keys match (handles minor variations)
                similar_structure = all(
                    isinstance(item, dict) and
                    len(set(item.keys()) & first_keys) >= len(first_keys) * 0.5
                    for item in data[1:]
                )
                if similar_structure:
                    # Collect all unique keys maintaining order
                    all_keys = list(data[0].keys())
                    for item in data[1:]:
                        for key in item.keys():
                            if key not in all_keys:
                                all_keys.append(key)
                    return format_table(data, headers=all_keys)

        # Otherwise format as nested structure
        result = ""
        for i, item in enumerate(data):
            if isinstance(item, dict):
                indent = "  " * indent_level
                # Add separator for multiple dict items
                if i > 0:
                    result += "\n"
                result += format_nested_dict(item, indent_level + 1)
            else:
                indent = "  " * indent_level
                result += f"{indent}- {safe_str(item)}\n"
        return result + "\n"

    if isinstance(data, dict):
        result = ""
        for key, value in data.items():
            indent = "  " * indent_level

            # Format key nicely
            formatted_key = safe_str(key).strip(':')

            if isinstance(value, dict):
                result += f"{indent}**{formatted_key}:**\n\n"
                result += format_nested_dict(value, indent_level + 1)
            elif isinstance(value, list):
                result += f"{indent}**{formatted_key}:**\n\n"

                # Check if it's a list of dicts (potential table)
                if value and isinstance(value[0], dict) and len(value) > 1:
                    # Check if all items are dicts (allow table with missing keys)
                    if all(isinstance(item, dict) for item in value):
                        result += format_table(value)
                        continue

                # Otherwise format as list
                for item in value:
                    if isinstance(item, dict):
                        result += format_nested_dict(item, indent_level + 1)
                    else:
                        result += f"{indent}- {safe_str(item)}\n"
                result += "\n"
            else:
                result += f"{indent}**{formatted_key}:** {safe_str(value)}\n\n"

        return result

    return f"{safe_str(data)}\n\n"


def format_findings_section(findings):
    """Format empirical findings with proper handling of complex data structures."""
    if findings is None:
        return "No findings available.\n"

    # Handle string findings
    if isinstance(findings, str):
        return f"{findings}\n"

    # Handle list of simple items
    if isinstance(findings, list):
        # Check if all items are simple strings/numbers
        if all(isinstance(item, (str, int, float, bool)) for item in findings):
            result = ""
            for finding in findings:
                result += f"- {safe_str(finding)}\n"
            return result + "\n"

        # Otherwise, complex list - format each item
        result = ""
        for i, item in enumerate(findings, 1):
            if isinstance(item, dict):
                result += format_nested_dict(item)
            else:
                result += f"{i}. {safe_str(item)}\n"
        return result + "\n"

    # Handle dictionary findings
    if isinstance(findings, dict):
        return format_nested_dict(findings)

    # Fallback
    return f"{safe_str(findings)}\n"


def format_analysis_as_markdown(analysis_data):
    """Format parsed analysis data as beautiful markdown."""
    if isinstance(analysis_data, str):
        # Try to parse if it's a string
        parsed = parse_json_from_response(analysis_data)
        if parsed:
            analysis_data = parsed
        else:
            # Return as-is if can't parse
            return analysis_data

    markdown = ""

    # 1. Citation
    if "1. Full Citation (APA 7th)" in analysis_data:
        markdown += f"## 📚 Full Citation\n\n"
        markdown += f"{analysis_data['1. Full Citation (APA 7th)']}\n\n"
        markdown += "---\n\n"

    # 2. Research Question & Hypotheses
    if "2. Core Research Question & Hypothesis(es)" in analysis_data:
        section = analysis_data["2. Core Research Question & Hypothesis(es)"]
        markdown += f"## 🎯 Research Question & Hypotheses\n\n"

        if isinstance(section, dict):
            if "Primary Question" in section:
                markdown += f"### Primary Research Question\n\n"
                markdown += f"{section['Primary Question']}\n\n"

            if "Hypotheses" in section and section["Hypotheses"]:
                markdown += f"### Hypotheses\n\n"
                hypotheses = section["Hypotheses"]
                if isinstance(hypotheses, list):
                    for i, hypothesis in enumerate(hypotheses, 1):
                        markdown += f"{i}. {hypothesis}\n"
                else:
                    markdown += f"{hypotheses}\n"
                markdown += "\n"
        else:
            markdown += f"{section}\n\n"

        markdown += "---\n\n"

    # 3. Theoretical Framework
    if "3. Theoretical Framework" in analysis_data:
        markdown += f"## 🧠 Theoretical Framework\n\n"
        markdown += f"{analysis_data['3. Theoretical Framework']}\n\n"
        markdown += "---\n\n"

    # 4. Methodology & Design
    if "4. Methodology & Design" in analysis_data:
        section = analysis_data["4. Methodology & Design"]
        markdown += f"## 🔬 Methodology & Design\n\n"

        if isinstance(section, dict):
            for key, value in section.items():
                markdown += f"**{key}:** "
                if isinstance(value, list):
                    markdown += "\n"
                    for item in value:
                        markdown += f"- {item}\n"
                else:
                    markdown += f"{value}\n"
                markdown += "\n"
        else:
            markdown += f"{section}\n\n"

        markdown += "---\n\n"

    # 5. Empirical Findings
    if "5. Empirical Findings" in analysis_data:
        findings = analysis_data["5. Empirical Findings"]
        markdown += f"## 📊 Empirical Findings\n\n"

        markdown += format_findings_section(findings)
        markdown += "\n---\n\n"

    # 6. Authors' Conclusions
    conclusions_key = "6. Authors' Stated Conclusions"
    if conclusions_key in analysis_data:
        markdown += "## 💡 Authors' Conclusions\n\n"
        markdown += f"{analysis_data[conclusions_key]}\n\n"
        markdown += "---\n\n"

    # 7. Limitations
    limitations_key = "7. Authors' Stated Limitations"
    if limitations_key in analysis_data:
        limitations = analysis_data[limitations_key]
        markdown += "## ⚠️ Limitations\n\n"

        if isinstance(limitations, list):
            for limitation in limitations:
                markdown += f"- {limitation}\n"
        else:
            markdown += f"{limitations}\n"

        markdown += "\n---\n\n"

    # 8. Critical Appraisal
    if "8. [MY ANALYSIS] Critical Appraisal & Integration" in analysis_data:
        section = analysis_data["8. [MY ANALYSIS] Critical Appraisal & Integration"]
        markdown += f"## 🔍 Critical Appraisal & Integration\n\n"

        if isinstance(section, dict):
            for key, value in section.items():
                markdown += f"### {key}\n\n"

                # Check if value is a list (for Connections and Unanswered Questions)
                if isinstance(value, list):
                    for item in value:
                        markdown += f"- {item}\n"
                    markdown += "\n"
                else:
                    markdown += f"{value}\n\n"
        else:
            markdown += f"{section}\n\n"

        markdown += "---\n\n"

    # 9. Attributes and Tags
    if "9. Attributes and tags" in analysis_data:
        section = analysis_data["9. Attributes and tags"]
        markdown += f"## 🏷️ Metadata & Tags\n\n"

        if isinstance(section, dict):
            # Basic attributes
            if "type:" in section:
                markdown += f"**Type:** {section['type:']}\n\n"
            if "year:" in section:
                markdown += f"**Year:** {section['year:']}\n\n"
            if "rating:" in section:
                rating = "⭐" * int(section['rating:'])
                markdown += f"**Rating:** {rating} ({section['rating:']}/5)\n\n"
            if "journal:" in section:
                markdown += f"**Journal:** {section['journal:']}\n\n"

            # Authors
            if "authors:" in section and isinstance(section["authors:"], list):
                markdown += f"**Authors:** "
                markdown += ", ".join(section["authors:"]) + "\n\n"

            # Topics
            if "topic/" in section and isinstance(section["topic/"], list):
                markdown += f"**Topics:** "
                markdown += " ".join(section["topic/"]) + "\n\n"

            # Methods
            if "method/" in section and isinstance(section["method/"], list):
                markdown += f"**Methods:** "
                markdown += " ".join(section["method/"]) + "\n\n"

            # Theory
            if "theory/" in section and isinstance(section["theory/"], list):
                markdown += f"**Theory:** "
                markdown += " ".join(section["theory/"]) + "\n\n"

            # Population
            if "population/" in section and isinstance(section["population/"], list):
                markdown += f"**Population:** "
                markdown += " ".join(section["population/"]) + "\n\n"
        else:
            markdown += f"{section}\n\n"

    return markdown


def create_markdown_from_analysis(analysis_text):
    """Convert analysis to markdown format with proper formatting."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Try to parse and format the analysis
    formatted = format_analysis_as_markdown(analysis_text)

    # Create final markdown with header
    markdown = f"# 📖 Literature Analysis Report\n\n"
    markdown += f"**Generated:** {timestamp}\n\n"
    markdown += f"---\n\n"
    markdown += formatted
    markdown += "\n---\n\n"
    markdown += "*Generated by Literature Assistant powered by AI*\n"

    return markdown


def create_docx_from_markdown(markdown_content, filename):
    """Create a DOCX file from markdown content with improved formatting."""
    import re

    try:
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font # type: ignore
        font.name = 'Calibri'
        font.size = Pt(11)

        # Parse markdown and add to document
        lines = markdown_content.split('\n')
        skip_next = False

        for i, line in enumerate(lines):
            if skip_next:
                skip_next = False
                continue

            original_line = line
            line = line.strip()

            if not line:
                continue

            # Handle headers (remove emojis for cleaner DOCX)
            if line.startswith('# '):
                text = re.sub(r'[📖🎯🧠🔬📊💡⚠️🔍🏷️]', '', line[2:]).strip()
                doc.add_heading(text, level=1)
            elif line.startswith('## '):
                text = re.sub(r'[📖🎯🧠🔬📊💡⚠️🔍🏷️📚]', '', line[3:]).strip()
                doc.add_heading(text, level=2)
            elif line.startswith('### '):
                text = line[4:].strip()
                doc.add_heading(text, level=3)
            elif line.startswith('#### '):
                text = line[5:].strip()
                doc.add_heading(text, level=4)
            # Handle horizontal rules
            elif line.startswith('---'):
                p = doc.add_paragraph()
                p.add_run('_' * 80)
            # Handle bullet lists
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                paragraph = doc.add_paragraph(style='List Bullet')
                add_formatted_text(paragraph, text)
            # Handle numbered lists
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s+', '', line)
                paragraph = doc.add_paragraph(style='List Number')
                add_formatted_text(paragraph, text)
            # Handle italic emphasis at start
            elif line.startswith('*') and not line.startswith('**'):
                paragraph = doc.add_paragraph()
                paragraph.add_run(line.strip('*')).italic = True
            # Regular paragraphs with possible bold text
            else:
                paragraph = doc.add_paragraph()
                add_formatted_text(paragraph, line)

        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return file_stream

    except Exception as error:
        raise Exception(f"Failed to create DOCX: {str(error)}")


def add_formatted_text(paragraph, text):
    """Add text to paragraph with bold/italic formatting."""
    import re

    # Handle **bold** text
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            paragraph.add_run(part.strip('*')).bold = True
        elif part:
            # Regular text
            paragraph.add_run(part)


@app.route('/')
def index():
    """Serve the frontend application."""
    return render_template('index.html')


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint."""
    return jsonify({
        'status': 'healthy',
        'openai_configured': client_openai is not None
    })


@app.route('/api/analyze', methods=['POST'])
def analyze_document():
    """Main endpoint to analyze PDF documents."""
    try:
        # Validate file upload
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not is_allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Only PDF files are allowed.'}), 400

        # Save uploaded file
        filename = secure_filename(file.filename) # type: ignore
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_filename = f"{timestamp}_{filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)

        try:
            # Extract text from PDF
            text_content = extract_text_from_pdf(file_path)

            # Load prompt template
            prompt_template = load_prompt_template()

            # Analyze with OpenAI
            analysis_result = analyze_with_openai(text_content, prompt_template)

            # Create markdown output
            markdown_output = create_markdown_from_analysis(analysis_result)

            # Clean up uploaded file
            os.remove(file_path)

            return jsonify({
                'success': True,
                'markdown': markdown_output,
                'provider': 'openai',
                'timestamp': timestamp
            })

        except Exception as processing_error:
            # Clean up on error
            if os.path.exists(file_path):
                os.remove(file_path)
            raise processing_error

    except Exception as error:
        return jsonify({'error': str(error)}), 500


@app.route('/api/download/<format>', methods=['POST'])
def download_file(format):
    """Endpoint to download analysis results."""
    try:
        data = request.get_json()
        content = data.get('content')
        filename = data.get('filename', 'analysis')

        if not content:
            return jsonify({'error': 'No content provided'}), 400

        if format == 'markdown':
            # Return markdown file
            file_stream = io.BytesIO(content.encode('utf-8'))
            return send_file(
                file_stream,
                mimetype='text/markdown',
                as_attachment=True,
                download_name=filename
            )

        elif format == 'docx':
            # Create and return DOCX file
            docx_stream = create_docx_from_markdown(content, filename)
            return send_file(
                docx_stream,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=filename
            )

        else:
            return jsonify({'error': 'Invalid format'}), 400

    except Exception as error:
        return jsonify({'error': str(error)}), 500


@app.errorhandler(413)
def file_too_large(error):
    """Handle file too large error."""
    return jsonify({'error': 'File size exceeds 10MB limit'}), 413


@app.errorhandler(500)
def internal_error(error):
    """Handle internal server errors."""
    return jsonify({'error': 'Internal server error'}), 500


if __name__ == '__main__':
    # Get port from environment variable (for Docker/Render) or use 5001
    port = int(os.getenv('PORT', 5001))
    # Use 0.0.0.0 for Docker, allows external connections
    host = os.getenv('HOST', '0.0.0.0')
    debug = os.getenv('FLASK_DEBUG', 'False').lower() == 'true'

    print(f"Starting server on {host}:{port}")
    app.run(debug=debug, port=port, host=host)
